from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.schemas.extracted_document import ExtractedDocument


class DocumentExtractor:
    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".txt",
        ".md",
        ".docx",
    }

    def extract(self, file_path: str | Path) -> ExtractedDocument:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Document not found: {path}"
            )

        extension = path.suffix.lower()

        if extension == ".pdf":
            return self._extract_pdf(path)

        if extension in {".txt", ".md"}:
            return self._extract_text(path)

        if extension == ".docx":
            return self._extract_docx(path)

        raise ValueError(
            f"Unsupported file type: {extension}"
        )

    def _extract_pdf(
        self,
        path: Path,
    ) -> ExtractedDocument:
        reader = PdfReader(path)

        pages: list[str] = []

        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text.strip())

        text = "\n\n".join(
            page
            for page in pages
            if page
        ).strip()

        return ExtractedDocument(
            text=text,
            page_count=len(reader.pages),
            metadata={
                "filename": path.name,
                "file_type": "pdf",
            },
        )

    def _extract_text(
        self,
        path: Path,
    ) -> ExtractedDocument:
        text = path.read_text(
            encoding="utf-8",
            errors="ignore",
        ).strip()

        return ExtractedDocument(
            text=text,
            page_count=1,
            metadata={
                "filename": path.name,
                "file_type": path.suffix.lower()[1:],
            },
        )

    def _extract_docx(
        self,
        path: Path,
    ) -> ExtractedDocument:
        document = DocxDocument(path)

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        table_texts: list[str] = []

        for table in document.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if cells:
                    table_texts.append(" | ".join(cells))

        sections = paragraphs + table_texts
        text = "\n\n".join(sections).strip()

        return ExtractedDocument(
            text=text,
            page_count=1,
            metadata={
                "filename": path.name,
                "file_type": "docx",
            },
        )